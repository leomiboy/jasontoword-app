import json
import io
from flask import Flask, render_template, request, send_file
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH

app = Flask(__name__)

def set_font(run, font_name='微軟正黑體', size=11, bold=False):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(size)
    run.font.bold = bold

def generate_docx(data):
    doc = Document()
    
    # 1. 取得文件總標題
    doc_title = data.get('document_title', '文件提取結果')
    
    # 2. 處理每一頁資料
    for i, page in enumerate(data.get('pages', [])):
        # 加入該頁的章節標題 (Section Title)
        section_title = page.get('section_title', doc_title)
        st_p = doc.add_paragraph()
        st_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        st_run = st_p.add_run(section_title)
        set_font(st_run, size=16, bold=True)

        # 加入頁碼標示
        page_label = page.get('page_label', '')
        pl_p = doc.add_paragraph()
        pl_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        pl_run = pl_p.add_run(f"頁面：{page_label}")
        set_font(pl_run, size=9, bold=False)

        # 3. 動態建立表格
        headers = page.get('headers', [])
        if headers:
            num_cols = len(headers)
            table = doc.add_table(rows=1, cols=num_cols)
            table.style = 'Table Grid'
            table.autofit = False # 關閉自動調整，手動設定欄寬

            # 設定標頭樣式與寬度
            hdr_cells = table.rows[0].cells
            for idx, h_text in enumerate(headers):
                # 針對常見的 3 欄格式進行寬度優化
                if num_cols == 3:
                    if idx == 0: hdr_cells[idx].width = Inches(0.8)  # 用字欄
                    elif idx == 1: hdr_cells[idx].width = Inches(4.2) # 例句欄 (加寬)
                    elif idx == 2: hdr_cells[idx].width = Inches(1.5) # 字義欄
                
                run = hdr_cells[idx].paragraphs[0].add_run(h_text)
                set_font(run, size=11, bold=True)
                hdr_cells[idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
            # 4. 填入資料內容
            for row_data in page.get('data', []):
                row_cells = table.add_row().cells
                for col_idx, cell_value in enumerate(row_data):
                    # 內容排版設定
                    cell_p = row_cells[col_idx].paragraphs[0]
                    
                    # 第一欄(關鍵字)與最後一欄(字義)置中，中間(例句)靠左
                    if num_cols == 3:
                        if col_idx == 1: cell_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        else: cell_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    else:
                        cell_p.alignment = WD_ALIGN_PARAGRAPH.LEFT

                    run = cell_p.add_run(str(cell_value))
                    # 讓第一欄文字加粗
                    set_font(run, size=10, bold=(col_idx == 0))

        # 若不是最後一頁，加入分頁符號
        if i < len(data['pages']) - 1:
            doc.add_page_break()

    # 將檔案存入記憶體流
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream, doc_title

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/convert', methods=['POST'])
def convert():
    if 'json_file' not in request.files:
        return "未選擇檔案", 400
    
    file = request.files['json_file']
    try:
        # 讀取上傳的內容
        content = file.read().decode('utf-8')
        data = json.loads(content)
        
        # 生成 Word
        docx_file, doc_title = generate_docx(data)
        
        # 動態設定下載檔名
        safe_filename = doc_title.replace(" ", "_")
        return send_file(
            docx_file,
            as_attachment=True,
            download_name=f"{safe_filename}.docx",
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
    except Exception as e:
        return f"發生錯誤: {str(e)}", 500

if __name__ == '__main__':
    app.run(debug=True)